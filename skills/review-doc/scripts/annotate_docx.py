#!/usr/bin/env python3
"""annotate_docx.py — Add comments or tracked changes to a Word document.

Takes a docx and a JSON annotations file, outputs annotated docx.

JSON format:
{
  "comments": [
    {"paragraph": 0, "text": "风险提示：此条款对乙方不利", "author": "手工川"},
    {"paragraph": 3, "start": 10, "end": 25, "text": "建议修改此处表述", "author": "手工川"}
  ],
  "revisions": [
    {"paragraph": 0, "old": "原文字", "new": "修改后文字", "author": "手工川"}
  ]
}

- comments[].paragraph: 0-based paragraph index
- comments[].start/end: optional character offsets within the paragraph to highlight
- comments[].text: the comment content
- revisions[].paragraph: 0-based paragraph index
- revisions[].old: text to find in that paragraph
- revisions[].new: replacement text (tracked as a revision)
"""

import argparse
import copy
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

NSMAP = {
    "w": W_NS,
    "r": R_NS,
    "w14": W14_NS,
}


def qn(tag: str) -> str:
    """Expand namespace prefix, e.g. 'w:comment' -> '{http://...}comment'."""
    prefix, local = tag.split(":")
    return f"{{{NSMAP[prefix]}}}{local}"


def _ensure_comments_part(doc: Document):
    """Ensure the document has a comments.xml part; create if missing."""
    from docx.opc.packuri import PackURI

    COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    # Check if comments part already exists
    doc_part = doc.part
    for rel in doc_part.rels.values():
        if rel.reltype == COMMENTS_REL:
            return rel.target_part

    # Create via python-docx's CommentsPart if available, else raw Part
    try:
        from docx.parts.comments import CommentsPart
        comments_part = CommentsPart.default(doc_part.package)
    except (ImportError, AttributeError):
        from docx.opc.part import Part
        COMMENTS_URI = "/word/comments.xml"
        COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "</w:comments>"
        )
        comments_part = Part(
            PackURI(COMMENTS_URI),
            COMMENTS_CT,
            comments_xml.encode("utf-8"),
            doc_part.package,
        )
    doc_part.relate_to(comments_part, COMMENTS_REL)
    return comments_part


def _add_comment_to_part(comments_part, comment_id: int, author: str, text: str, date: str):
    """Append a <w:comment> element to the comments.xml part."""
    # For CommentsPart (XmlPart subclass), modify _element directly.
    # XmlPart.blob serializes from _element, so _blob assignment is ignored.
    root = comments_part._element
    comment_el = etree.SubElement(root, qn("w:comment"))
    comment_el.set(qn("w:id"), str(comment_id))
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), date)

    p_el = etree.SubElement(comment_el, qn("w:p"))
    r_el = etree.SubElement(p_el, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _wrap_paragraph_with_comment(paragraph, comment_id: int, start: int = None, end: int = None):
    """Insert commentRangeStart/End markers around a paragraph or text range."""
    p_elem = paragraph._element

    range_start = etree.Element(qn("w:commentRangeStart"))
    range_start.set(qn("w:id"), str(comment_id))

    range_end = etree.Element(qn("w:commentRangeEnd"))
    range_end.set(qn("w:id"), str(comment_id))

    # Comment reference run (the small superscript marker)
    ref_run = etree.Element(qn("w:r"))
    rpr = etree.SubElement(ref_run, qn("w:rPr"))
    rstyle = etree.SubElement(rpr, qn("w:rStyle"))
    rstyle.set(qn("w:val"), "CommentReference")
    comment_ref = etree.SubElement(ref_run, qn("w:commentReference"))
    comment_ref.set(qn("w:id"), str(comment_id))

    if start is not None and end is not None:
        # Wrap specific character range — we need to split runs
        _insert_markers_at_offsets(p_elem, comment_id, start, end, range_start, range_end, ref_run)
    else:
        # Wrap entire paragraph.
        # If all text lives inside <w:ins> elements (tracked changes), place
        # markers *inside* those elements so Word can anchor the comment to
        # visible text.
        ins_elems = p_elem.findall(qn("w:ins"))
        has_regular_runs = any(
            r.find(qn("w:t")) is not None
            for r in p_elem.findall(qn("w:r"))
        )
        if ins_elems and not has_regular_runs:
            # All text is in <w:ins> — anchor inside first/last ins
            first_ins = ins_elems[0]
            last_ins = ins_elems[-1]
            first_ins.insert(0, range_start)
            last_ins.append(range_end)
            last_ins.append(ref_run)
        else:
            p_elem.insert(0, range_start)
            p_elem.append(range_end)
            p_elem.append(ref_run)


def _insert_markers_at_offsets(p_elem, comment_id, start, end, range_start, range_end, ref_run):
    """Insert comment range markers at character offsets within a paragraph.

    Falls back to wrapping the whole paragraph if offset tracking fails.
    """
    runs = p_elem.findall(qn("w:r"))
    char_pos = 0
    start_inserted = False
    end_inserted = False

    for run in runs:
        t_elem = run.find(qn("w:t"))
        if t_elem is None or t_elem.text is None:
            continue
        run_start = char_pos
        run_end = char_pos + len(t_elem.text)

        if not start_inserted and run_start <= start < run_end:
            run.addprevious(range_start)
            start_inserted = True

        if not end_inserted and run_start < end <= run_end:
            run.addnext(range_end)
            range_end.addnext(ref_run)
            end_inserted = True

        char_pos = run_end

    # Fallback: wrap whole paragraph
    if not start_inserted:
        p_elem.insert(0, range_start)
    if not end_inserted:
        p_elem.append(range_end)
        p_elem.append(ref_run)


def _add_tracked_revision(paragraph, old_text: str, new_text: str, author: str, date: str, rev_id: int):
    """Add a tracked change (deletion + insertion) for a text substitution."""
    for run in paragraph.runs:
        if old_text in run.text:
            parts = run.text.split(old_text, 1)
            run_elem = run._element
            parent = run_elem.getparent()
            idx = list(parent).index(run_elem)

            # Preserve run properties
            rpr = run_elem.find(qn("w:rPr"))
            rpr_copy = copy.deepcopy(rpr) if rpr is not None else None

            # Before text (keep original run)
            if parts[0]:
                run.text = parts[0]
                idx += 1
            else:
                parent.remove(run_elem)

            # Deletion markup
            del_elem = etree.Element(qn("w:del"))
            del_elem.set(qn("w:id"), str(rev_id))
            del_elem.set(qn("w:author"), author)
            del_elem.set(qn("w:date"), date)
            del_run = etree.SubElement(del_elem, qn("w:r"))
            if rpr_copy is not None:
                del_run.append(copy.deepcopy(rpr_copy))
            del_text = etree.SubElement(del_run, qn("w:delText"))
            del_text.text = old_text
            del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            parent.insert(idx, del_elem)
            idx += 1

            # Insertion markup
            ins_elem = etree.Element(qn("w:ins"))
            ins_elem.set(qn("w:id"), str(rev_id + 1))
            ins_elem.set(qn("w:author"), author)
            ins_elem.set(qn("w:date"), date)
            ins_run = etree.SubElement(ins_elem, qn("w:r"))
            if rpr_copy is not None:
                ins_run.append(copy.deepcopy(rpr_copy))
            ins_text = etree.SubElement(ins_run, qn("w:t"))
            ins_text.text = new_text
            ins_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            parent.insert(idx, ins_elem)
            idx += 1

            # After text
            if parts[1]:
                after_run = etree.Element(qn("w:r"))
                if rpr_copy is not None:
                    after_run.append(copy.deepcopy(rpr_copy))
                after_t = etree.SubElement(after_run, qn("w:t"))
                after_t.text = parts[1]
                after_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                parent.insert(idx, after_run)

            return True

    return False


def annotate(input_path: str, annotations_path: str, output_path: str):
    """Apply annotations from JSON to a docx file."""
    doc = Document(input_path)
    with open(annotations_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%S+08:00")
    paragraphs = doc.paragraphs

    # Process comments
    comments = data.get("comments", [])
    if comments:
        comments_part = _ensure_comments_part(doc)

        # Find max existing comment ID to avoid collisions
        existing_root = comments_part._element
        existing_comments = existing_root.findall(qn("w:comment"))
        if existing_comments:
            next_comment_id = max(int(c.get(qn("w:id")) or "0") for c in existing_comments) + 1
        else:
            next_comment_id = 0

        # Also check body for commentRangeStart/End IDs
        body = doc.element.body
        for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
            for elem in body.iter(qn(tag)):
                eid = int(elem.get(qn("w:id")) or "0")
                if eid >= next_comment_id:
                    next_comment_id = eid + 1

        for i, c in enumerate(comments):
            para_idx = c["paragraph"]
            if para_idx >= len(paragraphs):
                print(f"WARN: paragraph index {para_idx} out of range ({len(paragraphs)} total), skipping", file=sys.stderr)
                continue
            comment_id = next_comment_id + i
            author = c.get("author", "Reviewer")
            _add_comment_to_part(comments_part, comment_id, author, c["text"], now)
            _wrap_paragraph_with_comment(
                paragraphs[para_idx], comment_id,
                start=c.get("start"), end=c.get("end"),
            )

    # Process revisions (tracked changes)
    revisions = data.get("revisions", [])
    # Find max ID used in the document to avoid collision
    rev_id = next_comment_id + len(comments) + 100 if comments else 1000
    for tag in ("w:ins", "w:del"):
        for elem in doc.element.body.iter(qn(tag)):
            eid = int(elem.get(qn("w:id")) or "0")
            if eid >= rev_id:
                rev_id = eid + 1
    for r in revisions:
        para_idx = r["paragraph"]
        if para_idx >= len(paragraphs):
            print(f"WARN: paragraph index {para_idx} out of range, skipping", file=sys.stderr)
            continue
        author = r.get("author", "Reviewer")
        ok = _add_tracked_revision(paragraphs[para_idx], r["old"], r["new"], author, now, rev_id)
        if not ok:
            print(f"WARN: could not find '{r['old']}' in paragraph {para_idx}, skipping", file=sys.stderr)
        rev_id += 2  # Each revision uses 2 IDs (del + ins)

    doc.save(output_path)
    total = len(comments) + len(revisions)
    print(f"Done! {len(comments)} comments, {len(revisions)} revisions → {output_path}")
    return total


def _get_full_paragraph_text(p_elem) -> str:
    """Get all text from a paragraph, including text inside <w:ins> tracked changes."""
    texts = []
    for elem in p_elem.iter(qn("w:t")):
        if elem.text:
            texts.append(elem.text)
    return "".join(texts).strip()


def extract_text(input_path: str):
    """Extract paragraph text with indices for AI review reference.

    Includes text inside <w:ins> (tracked changes) which python-docx's
    paragraph.text does not return.
    """
    doc = Document(input_path)
    result = []
    for i, p in enumerate(doc.paragraphs):
        # Try full text first (includes tracked changes), fall back to .text
        text = _get_full_paragraph_text(p._element)
        if text:
            result.append({"index": i, "text": text})
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)


def main():
    parser = argparse.ArgumentParser(description="Annotate a Word document with comments and tracked changes")
    sub = parser.add_subparsers(dest="command")

    # annotate subcommand
    ann = sub.add_parser("annotate", help="Apply annotations from JSON to docx")
    ann.add_argument("--input", required=True, help="Input docx path")
    ann.add_argument("--annotations", required=True, help="JSON annotations file")
    ann.add_argument("--output", required=True, help="Output docx path")

    # extract subcommand
    ext = sub.add_parser("extract", help="Extract paragraph text as JSON (for AI review)")
    ext.add_argument("--input", required=True, help="Input docx path")

    args = parser.parse_args()

    if args.command == "annotate":
        annotate(args.input, args.annotations, args.output)
    elif args.command == "extract":
        extract_text(args.input)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
