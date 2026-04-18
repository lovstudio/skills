#!/usr/bin/env python3
"""
fill_form — Fill in Word document form templates.

Reads a .docx template containing tables with label→value cell pairs,
fills specified fields with provided data, and saves the result.

Supports:
  - .doc input (auto-converts via textutil on macOS)
  - Table-based forms (label in one cell, value in adjacent cell)
  - Multi-row merged cells (e.g., large text areas)
  - CJK/Latin mixed text with font switching
  - Scan mode to list all detected fields

Usage:
  # Scan template to list fields
  python fill_form.py --template form.docx --scan

  # Fill form with JSON data
  python fill_form.py --template form.docx --output filled.docx \\
    --data '{"主讲人": "张三", "讲座题目": "AI与未来"}'

  # Fill from JSON file
  python fill_form.py --template form.docx --output filled.docx --data-file data.json

Dependencies:
  pip install python-docx --break-system-packages
"""

import re, os, sys, json, argparse, shutil, subprocess, tempfile
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import platform as _platform

_PLAT = _platform.system()

# ─── Fonts ───────────────────────────────────────────────────────────
def _get_fonts():
    if _PLAT == "Darwin":
        return ("Songti SC", "PingFang SC", "Menlo")
    elif _PLAT == "Windows":
        return ("SimSun", "Microsoft YaHei", "Consolas")
    else:
        return ("Noto Serif CJK SC", "Noto Sans CJK SC", "DejaVu Sans Mono")

CJK_SERIF, CJK_SANS, MONO = _get_fonts()

_CJK_RANGES = [
    (0x4E00, 0x9FFF), (0x3400, 0x4DBF), (0xF900, 0xFAFF), (0x3000, 0x303F),
    (0xFF00, 0xFFEF), (0x2E80, 0x2EFF), (0x2F00, 0x2FDF), (0xFE30, 0xFE4F),
    (0x20000, 0x2A6DF),
]

def _is_cjk(ch):
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in _CJK_RANGES)


# ─── .doc → .docx conversion ────────────────────────────────────────
def convert_doc_to_docx(doc_path):
    """Convert .doc to .docx, return path to temp .docx file."""
    tmp_dir = tempfile.mkdtemp(prefix="fillform_")
    base = os.path.splitext(os.path.basename(doc_path))[0]
    out_path = os.path.join(tmp_dir, base + ".docx")

    if _PLAT == "Darwin":
        # Try textutil first (preserves basic structure)
        r = subprocess.run(
            ["textutil", "-convert", "docx", "-output", out_path, doc_path],
            capture_output=True, text=True,
        )
        if r.returncode == 0 and os.path.exists(out_path):
            return out_path

    # Try LibreOffice as fallback
    for soffice in ["soffice", "libreoffice",
                     "/Applications/LibreOffice.app/Contents/MacOS/soffice"]:
        if shutil.which(soffice) or os.path.exists(soffice):
            r = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", tmp_dir, doc_path],
                capture_output=True, text=True,
            )
            candidate = os.path.join(tmp_dir, base + ".docx")
            if r.returncode == 0 and os.path.exists(candidate):
                return candidate

    print(f"ERROR: Cannot convert .doc to .docx. Install LibreOffice or use a .docx file.",
          file=sys.stderr)
    sys.exit(1)


# ─── Field detection ────────────────────────────────────────────────
def _normalize_label(text):
    """Normalize a label: strip whitespace and common padding characters."""
    return re.sub(r'[\s\u3000\xa0]+', '', text).strip()


def _cell_text(cell):
    """Get clean text from a cell."""
    return cell.text.strip()


def _is_label_cell(text):
    """Heuristic: a cell is a label if it has text, is short, and looks like a field name."""
    clean = _normalize_label(text)
    if not clean or len(clean) > 50:
        return False
    # Must contain at least one CJK char or look like a known pattern
    if any(_is_cjk(c) for c in clean):
        return True
    if re.match(r'^[A-Za-z\s/]+$', clean) and len(clean) < 30:
        return True
    return False


def scan_fields(doc):
    """Scan document tables and return list of detected field labels with locations."""
    fields = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = row.cells
            ci = 0
            while ci < len(cells):
                cell = cells[ci]
                text = _cell_text(cell)
                if _is_label_cell(text):
                    norm = _normalize_label(text)
                    # The value cell is the next non-label cell
                    value_ci = ci + 1
                    # Skip duplicate merged cells (python-docx repeats merged cells)
                    while value_ci < len(cells) and cells[value_ci]._tc is cell._tc:
                        value_ci += 1
                    if value_ci < len(cells):
                        value_cell = cells[value_ci]
                        existing = _cell_text(value_cell)
                        fields.append({
                            "label": norm,
                            "raw_label": text,
                            "table": ti,
                            "row": ri,
                            "label_col": ci,
                            "value_col": value_ci,
                            "current_value": existing if existing else "",
                        })
                        ci = value_ci + 1
                        continue
                ci += 1

    # Also detect large text areas (merged cells spanning full row with label in text)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = row.cells
            if len(set(id(c._tc) for c in cells)) == 1:
                # All cells are the same (fully merged row)
                text = _cell_text(cells[0])
                # Check if it starts with a label-like pattern
                m = re.match(r'^(.+?)[：:]\s*$', text) or re.match(r'^(.+?)[：:]', text)
                if m:
                    label = _normalize_label(m.group(1))
                    # Avoid duplicates
                    if not any(f["label"] == label for f in fields):
                        fields.append({
                            "label": label,
                            "raw_label": m.group(1),
                            "table": ti,
                            "row": ri,
                            "label_col": 0,
                            "value_col": -1,  # -1 means inline (same cell, after colon)
                            "current_value": text[m.end():].strip(),
                        })

    # Fallback: if no table fields found, try paragraph-based detection
    if not fields:
        fields = _scan_paragraph_fields(doc)

    return fields


def _scan_paragraph_fields(doc):
    """Detect fields in paragraph-based forms (no tables).
    Pattern: a paragraph with label text followed by an empty/short paragraph as value.
    Also detects 'label：value' patterns on a single line.
    """
    fields = []
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        text = p.text.strip()
        if not text:
            continue

        # Pattern 1: "Label：" at end of line, value on next line(s) or empty
        m = re.match(r'^(.+?)[：:]\s*$', text)
        if m and _is_label_cell(m.group(1)):
            label = _normalize_label(m.group(1))
            # Value is in following non-empty paragraph(s) until next label
            value_parts = []
            j = i + 1
            while j < len(paras):
                next_text = paras[j].text.strip()
                if next_text and not re.match(r'^.+?[：:]\s*$', next_text) and not _is_label_cell(next_text):
                    value_parts.append(next_text)
                else:
                    break
                j += 1
            fields.append({
                "label": label,
                "raw_label": m.group(1),
                "table": -1,
                "row": -1,
                "label_col": -1,
                "value_col": -1,
                "para_index": i,
                "current_value": " ".join(value_parts),
                "type": "paragraph",
            })
            continue

        # Pattern 2: "Label：value" on same line
        m = re.match(r'^(.+?)[：:]\s*(.+)$', text)
        if m and _is_label_cell(m.group(1)):
            label = _normalize_label(m.group(1))
            if not any(f["label"] == label for f in fields):
                fields.append({
                    "label": label,
                    "raw_label": m.group(1),
                    "table": -1,
                    "row": -1,
                    "label_col": -1,
                    "value_col": -1,
                    "para_index": i,
                    "current_value": m.group(2).strip(),
                    "type": "paragraph_inline",
                })

    return fields


def _set_cell_text(cell, text, font_name=CJK_SERIF, font_size=Pt(11)):
    """Set cell text, preserving basic formatting and applying CJK font."""
    # Clear existing content
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    # Preserve alignment if set
    run = p.add_run(text)
    run.font.size = font_size
    # Set both Latin and CJK font
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def fill_fields(doc, data, fields=None, font_name=CJK_SERIF, font_size=Pt(11)):
    """Fill detected fields with provided data. Returns list of filled field names."""
    if fields is None:
        fields = scan_fields(doc)

    filled = []
    args_font = font_name
    args_font_size = font_size
    data_norm = {_normalize_label(k): v for k, v in data.items()}

    for field in fields:
        label = field["label"]
        if label not in data_norm:
            continue

        value = str(data_norm[label])
        ti = field["table"]
        ri = field["row"]

        if ti >= 0:
            table = doc.tables[ti]
            row = table.rows[ri]
        else:
            row = None

        field_type = field.get("type", "table")

        if field_type == "paragraph":
            # Paragraph-based: set the paragraph after the label
            pi = field["para_index"]
            p = doc.paragraphs[pi]
            raw = field["raw_label"]
            # Clear following value paragraphs and set first one
            if pi + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[pi + 1]
                next_p.clear()
                run = next_p.add_run(value)
                run.font.name = args_font
                run.font.size = args_font_size
            else:
                p.clear()
                run = p.add_run(f"{raw}：{value}")
                run.font.name = args_font
                run.font.size = args_font_size
        elif field_type == "paragraph_inline":
            pi = field["para_index"]
            p = doc.paragraphs[pi]
            raw = field["raw_label"]
            p.clear()
            run = p.add_run(f"{raw}：{value}")
            run.font.name = args_font
            run.font.size = args_font_size
        elif row is not None and field["value_col"] == -1:
            # Inline field in table cell (label: value in same cell)
            cell = row.cells[field["label_col"]]
            raw = field["raw_label"]
            _set_cell_text(cell, f"{raw}：{value}")
        elif row is not None:
            cell = row.cells[field["value_col"]]
            _set_cell_text(cell, value)

        filled.append(label)

    return filled


# We need nsdecls import for _set_cell_text
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ─── Main ───────────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser(description="Fill in Word form templates")
    ap.add_argument("--template", required=True, help="Path to template .doc/.docx file")
    ap.add_argument("--output", default="", help="Output .docx path (default: <template>_filled.docx)")
    ap.add_argument("--scan", action="store_true", help="Scan and list all detected form fields")
    ap.add_argument("--data", default="", help="JSON string with field→value mapping")
    ap.add_argument("--data-file", default="", help="Path to JSON file with field→value mapping")
    ap.add_argument("--font", default=CJK_SERIF, help=f"Font name for filled text (default: {CJK_SERIF})")
    ap.add_argument("--font-size", type=float, default=11, help="Font size in points (default: 11)")
    args = ap.parse_args()

    template_path = args.template

    # Handle .doc files
    tmp_docx = None
    if template_path.lower().endswith(".doc") and not template_path.lower().endswith(".docx"):
        print(f"Converting .doc → .docx ...", file=sys.stderr)
        tmp_docx = convert_doc_to_docx(template_path)
        template_path = tmp_docx

    doc = Document(template_path)

    if args.scan:
        fields = scan_fields(doc)
        if not fields:
            print("No form fields detected in the document.")
            return
        print(f"Detected {len(fields)} form fields:\n")
        for i, f in enumerate(fields, 1):
            current = f" (current: \"{f['current_value']}\")" if f["current_value"] else ""
            print(f"  {i}. {f['label']}{current}")
        # Also output as JSON for programmatic use
        print(f"\n--- JSON ---")
        template = {f["label"]: f["current_value"] or "" for f in fields}
        print(json.dumps(template, ensure_ascii=False, indent=2))
        return

    # Load data
    if args.data:
        data = json.loads(args.data)
    elif args.data_file:
        with open(args.data_file, encoding="utf-8") as f:
            data = json.load(f)
    else:
        print("ERROR: Provide --data or --data-file to fill the form.", file=sys.stderr)
        sys.exit(1)

    fields = scan_fields(doc)
    filled = fill_fields(doc, data, fields, font_name=args.font, font_size=Pt(args.font_size))

    # Determine output path
    output = args.output
    if not output:
        base = os.path.splitext(os.path.basename(args.template))[0]
        template_dir = os.path.dirname(os.path.abspath(args.template))
        output = os.path.join(template_dir, base + "_filled.docx")

    doc.save(output)
    print(f"Filled {len(filled)}/{len(fields)} fields → {output}")
    if filled:
        print(f"  Filled: {', '.join(filled)}")
    unfilled = [f["label"] for f in fields if f["label"] not in filled]
    if unfilled:
        print(f"  Unfilled: {', '.join(unfilled)}")


if __name__ == "__main__":
    main()
