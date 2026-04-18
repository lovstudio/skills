#!/usr/bin/env python3
"""
md2docx — Convert Markdown to professionally styled DOCX.

Features:
  - CJK/Latin mixed text with automatic font switching
  - Fenced code blocks with shaded background
  - Markdown tables with header row styling
  - Cover page, table of contents
  - Configurable color themes (same palette as any2pdf)
  - Watermark support
  - Headers and footers with page numbers

Usage:
  python md2docx.py --input report.md --output report.docx --title "My Report"

Dependencies:
  pip install python-docx --break-system-packages
"""

import re, os, sys, argparse, tempfile, urllib.request
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import platform as _platform

_PLAT = _platform.system()

# ═══════════════════════════════════════════════════════════════════════
# FONTS — platform-aware CJK font names for DOCX
# ═══════════════════════════════════════════════════════════════════════
def _get_fonts():
    """Return (serif, sans, mono, cjk) font names available on this platform."""
    if _PLAT == "Darwin":
        return ("Palatino", "Arial", "Menlo", "Songti SC")
    elif _PLAT == "Windows":
        return ("Times New Roman", "Arial", "Consolas", "SimSun")
    else:  # Linux
        return ("Liberation Serif", "Liberation Sans", "DejaVu Sans Mono", "Noto Serif CJK SC")

SERIF, SANS, MONO, CJK = _get_fonts()

# ═══════════════════════════════════════════════════════════════════════
# CJK DETECTION
# ═══════════════════════════════════════════════════════════════════════
_CJK_RANGES = [
    (0x4E00,0x9FFF),(0x3400,0x4DBF),(0xF900,0xFAFF),(0x3000,0x303F),
    (0xFF00,0xFFEF),(0x2E80,0x2EFF),(0x2F00,0x2FDF),(0xFE30,0xFE4F),
    (0x20000,0x2A6DF),(0x2A700,0x2B73F),(0x2B740,0x2B81F),
]

def _is_cjk(ch):
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in _CJK_RANGES)

def _split_mixed(text):
    """Split text into segments of (is_cjk, text) for font switching."""
    if not text:
        return []
    segs, buf, in_cjk = [], [], False
    for ch in text:
        c = _is_cjk(ch)
        if c != in_cjk and buf:
            segs.append((in_cjk, ''.join(buf)))
            buf = []
        buf.append(ch)
        in_cjk = c
    if buf:
        segs.append((in_cjk, ''.join(buf)))
    return segs

# ═══════════════════════════════════════════════════════════════════════
# THEMES — same palette as any2pdf
# ═══════════════════════════════════════════════════════════════════════
THEMES = {
    "warm-academic": {
        "ink":"#181818","ink_faded":"#87867F","accent":"#CC785C",
        "border":"#E8E6DC","canvas_sec":"#F0EEE6",
        "body_font":"serif","heading_font":"serif",
    },
    "nord-frost": {
        "ink":"#2E3440","ink_faded":"#4C566A","accent":"#5E81AC",
        "border":"#D8DEE9","canvas_sec":"#E5E9F0",
        "body_font":"sans","heading_font":"sans",
    },
    "github-light": {
        "ink":"#1F2328","ink_faded":"#656D76","accent":"#0969DA",
        "border":"#D0D7DE","canvas_sec":"#F6F8FA",
        "body_font":"sans","heading_font":"sans",
    },
    "solarized-light": {
        "ink":"#657B83","ink_faded":"#93A1A1","accent":"#CB4B16",
        "border":"#EEE8D5","canvas_sec":"#EEE8D5",
        "body_font":"serif","heading_font":"serif",
    },
    "ocean-breeze": {
        "ink":"#1A2E35","ink_faded":"#5A7D7C","accent":"#2A9D8F",
        "border":"#C8DDD6","canvas_sec":"#E0EDE8",
        "body_font":"sans","heading_font":"sans",
    },
    "tufte": {
        "ink":"#111111","ink_faded":"#999988","accent":"#980000",
        "border":"#E0DDD0","canvas_sec":"#F7F7F0",
        "body_font":"serif","heading_font":"serif",
    },
    "classic-thesis": {
        "ink":"#2B2B2B","ink_faded":"#7A7568","accent":"#8B4513",
        "border":"#D6CFC2","canvas_sec":"#F5F2EB",
        "body_font":"serif","heading_font":"serif",
    },
    "ieee-journal": {
        "ink":"#000000","ink_faded":"#555555","accent":"#003366",
        "border":"#CCCCCC","canvas_sec":"#F5F5F5",
        "body_font":"serif","heading_font":"serif",
    },
    "elegant-book": {
        "ink":"#1A1A1A","ink_faded":"#6E6B5E","accent":"#5B3A29",
        "border":"#DDD8C8","canvas_sec":"#F0ECE0",
        "body_font":"serif","heading_font":"serif",
    },
    "chinese-red": {
        "ink":"#1A1009","ink_faded":"#8C7A5E","accent":"#B22222",
        "border":"#E8DCC8","canvas_sec":"#F8F0E0",
        "body_font":"serif","heading_font":"serif",
    },
    "ink-wash": {
        "ink":"#2C2C2C","ink_faded":"#8A8A80","accent":"#404040",
        "border":"#D8D4C8","canvas_sec":"#EEEAE0",
        "body_font":"serif","heading_font":"serif",
    },
    "invest-report": {
        "ink":"#1A1A1A","ink_faded":"#666666","accent":"#C00000",
        "border":"#CCCCCC","canvas_sec":"#F5F5F5",
        "body_font":"serif","heading_font":"serif",
        "cjk_font":"STKaiti",
    },
}

def _hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def load_theme(name):
    t = THEMES.get(name)
    if not t:
        print(f"Unknown theme '{name}', falling back to warm-academic", file=sys.stderr)
        t = THEMES["warm-academic"]
    result = {
        "ink": _hex_to_rgb(t["ink"]),
        "ink_faded": _hex_to_rgb(t["ink_faded"]),
        "accent": _hex_to_rgb(t["accent"]),
        "border": _hex_to_rgb(t["border"]),
        "canvas_sec": t["canvas_sec"],
        "body_font": SERIF if t["body_font"] == "serif" else SANS,
        "heading_font": SERIF if t["heading_font"] == "serif" else SANS,
    }
    if "cjk_font" in t:
        result["cjk_font"] = t["cjk_font"]
    return result

# ═══════════════════════════════════════════════════════════════════════
# INLINE MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════════════════
def _parse_inline(text):
    """Parse inline markdown into segments: (text, bold, italic, code, link_url)."""
    segments = []
    # Tokenize: **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'        # bold
        r'|(`(.+?)`)'              # code
        r'|(\*(.+?)\*)'            # italic
        r'|(\[(.+?)\]\((.+?)\))'   # link
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append({"text": text[pos:m.start()]})
        if m.group(2):  # bold
            segments.append({"text": m.group(2), "bold": True})
        elif m.group(4):  # code
            segments.append({"text": m.group(4), "code": True})
        elif m.group(6):  # italic
            segments.append({"text": m.group(6), "italic": True})
        elif m.group(8):  # link
            segments.append({"text": m.group(8), "link": m.group(9)})
        pos = m.end()
    if pos < len(text):
        segments.append({"text": text[pos:]})
    return segments

def _add_mixed_run(paragraph, text, font_name, font_size, color=None,
                   bold=False, italic=False):
    """Add text with CJK font switching as multiple runs."""
    for is_cjk, seg in _split_mixed(text):
        run = paragraph.add_run(seg)
        run.font.name = CJK if is_cjk else font_name
        # Set East Asian font for CJK
        run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
        run.font.size = font_size
        if color:
            run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic

def _add_inline_content(paragraph, text, theme, font_name, font_size):
    """Add inline-markdown-parsed content to a paragraph with CJK support."""
    segments = _parse_inline(text)
    for seg in segments:
        t = seg["text"]
        if seg.get("code"):
            for is_cjk, part in _split_mixed(t):
                run = paragraph.add_run(part)
                run.font.name = CJK if is_cjk else MONO
                run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
                run.font.size = Pt(font_size.pt - 1 if hasattr(font_size, 'pt') else 9)
                run.font.color.rgb = theme["accent"]
        else:
            _add_mixed_run(paragraph, t, font_name, font_size,
                          color=theme["ink"],
                          bold=seg.get("bold", False),
                          italic=seg.get("italic", False))

# ═══════════════════════════════════════════════════════════════════════
# DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════
class DocxBuilder:
    def __init__(self, config):
        self.cfg = config
        self.T = config["theme"]
        # Override global CJK font if theme specifies one (e.g. STKaiti for invest-report)
        global CJK
        if "cjk_font" in self.T:
            CJK = self.T["cjk_font"]
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document default styles."""
        T = self.T
        style = self.doc.styles['Normal']
        style.font.name = T["body_font"]
        style.font.size = Pt(10.5)
        style.font.color.rgb = T["ink"]
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(4)
        # Set East Asian font
        style._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

        # Heading styles
        for level, size in [(1, 26), (2, 18), (3, 13)]:
            hs = self.doc.styles[f'Heading {level}']
            hs.font.name = T["heading_font"]
            hs.font.size = Pt(size)
            hs.font.color.rgb = T["accent"] if level == 3 else T["ink"]
            hs.font.bold = True
            hs._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
            hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            hs.paragraph_format.space_after = Pt(6)
            if level <= 2:
                hs.paragraph_format.page_break_before = True

        # Set page margins
        section = self.doc.sections[0]
        section.top_margin = Mm(28)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(22)

    def _add_watermark(self, text):
        """Add diagonal watermark to document header using raw XML with full namespace URIs."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        from lxml import etree
        NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        NS_V = 'urn:schemas-microsoft-com:vml'
        NS_O = 'urn:schemas-microsoft-com:office:office'
        wm_xml = (
            f'<w:r xmlns:w="{NS_W}" xmlns:v="{NS_V}" xmlns:o="{NS_O}">'
            f'<w:rPr><w:noProof/></w:rPr>'
            f'<w:pict>'
            f'<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" '
            f'path="m@7,l@8,m@5,21600l@6,21600e">'
            f'<v:formulas><v:f eqn="sum #0 0 10800"/></v:formulas>'
            f'<v:path textpathok="t" o:connecttype="custom"/>'
            f'<v:textpath on="t" fitshape="t"/>'
            f'<o:lock v:ext="edit" text="t" shapetype="t"/>'
            f'</v:shapetype>'
            f'<v:shape id="WaterMark" o:spid="_x0000_s2049" type="#_x0000_t136" '
            f'style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;'
            f'rotation:315;z-index:-251657216;'
            f'mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
            f'mso-position-vertical:center;mso-position-vertical-relative:margin" '
            f'o:allowincell="f" fillcolor="silver" stroked="f">'
            f'<v:fill opacity=".15"/>'
            f'<v:textpath style="font-family:&quot;Arial&quot;;font-size:1pt" string="{text}"/>'
            f'</v:shape>'
            f'</w:pict>'
            f'</w:r>'
        )
        p._element.append(etree.fromstring(wm_xml))

    def _add_cover(self):
        """Add cover page."""
        T = self.T
        # Add spacing before title
        for _ in range(6):
            self.doc.add_paragraph()

        # Title — adaptive font size based on length
        title = self.cfg.get("title", "Document")
        # CJK chars are ~2x width of Latin chars for layout estimation
        visual_len = sum(2 if _is_cjk(ch) else 1 for ch in title)
        if visual_len <= 16:
            title_size = 36
        elif visual_len <= 24:
            title_size = 30
        elif visual_len <= 32:
            title_size = 26
        else:
            title_size = 22
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_mixed_run(p, title, T["heading_font"], Pt(title_size), T["ink"], bold=True)

        # Subtitle
        sub = self.cfg.get("subtitle", "")
        if sub:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_mixed_run(p, sub, T["body_font"], Pt(16), T["ink_faded"])

        # Version
        ver = self.cfg.get("version", "")
        if ver:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_mixed_run(p, ver, SANS, Pt(12), T["accent"])

        # Separator
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("─" * 20)
        run.font.color.rgb = T["accent"]
        run.font.size = Pt(10)

        # Stats lines
        for key in ("stats_line", "stats_line2"):
            val = self.cfg.get(key, "")
            if val:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_mixed_run(p, val, SANS, Pt(9), T["ink_faded"])

        # Author + date at bottom
        for _ in range(4):
            self.doc.add_paragraph()

        author = self.cfg.get("author", "")
        if author:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_mixed_run(p, author, T["body_font"], Pt(11), T["ink_faded"])

        dt = self.cfg.get("date", str(date.today()))
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_mixed_run(p, dt, T["body_font"], Pt(10), T["ink_faded"])

        edition = self.cfg.get("edition_line", "")
        if edition:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_mixed_run(p, edition, T["body_font"], Pt(8), T["ink_faded"])

        # Page break after cover
        self.doc.add_page_break()

    def _add_toc(self, md_text=""):
        """Add TOC using field code + static fallback entries."""
        T = self.T
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_mixed_run(p, "目    录", T["heading_font"], Pt(26), T["ink"], bold=True)

        self.doc.add_paragraph()  # spacing

        # Insert TOC field code — auto-updated via updateFields setting
        p = self.doc.add_paragraph()
        run = p.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fld_char_begin)
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run2._element.append(instr)
        run3 = p.add_run()
        fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._element.append(fld_char_sep)

        # Static fallback: parse headings so TOC is visible even before refresh
        in_code = False
        for line in md_text.split('\n'):
            stripped = line.strip()
            if stripped.startswith('```'):
                in_code = not in_code
                continue
            if in_code:
                continue
            m = re.match(r'^(#{1,3})\s+(.+)$', stripped)
            if m:
                level = len(m.group(1))
                title = re.sub(r'\*\*(.+?)\*\*', r'\1', m.group(2).strip())
                title = re.sub(r'`(.+?)`', r'\1', title)
                prefix = "    " * (level - 1)
                run_entry = p.add_run(f"\n{prefix}{title}")
                run_entry.font.size = Pt(10.5 if level == 1 else 9.5)
                run_entry.font.color.rgb = T["ink"]
                run_entry.font.bold = (level <= 1)

        run5 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._element.append(fld_char_end)

        self.doc.add_page_break()

    def _enable_auto_update_fields(self):
        """Set updateFields=true so Word auto-refreshes TOC on open."""
        from lxml import etree
        settings = self.doc.settings.element
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Remove existing updateFields if any
        for el in settings.findall(f'{{{ns}}}updateFields'):
            settings.remove(el)
        uf = etree.SubElement(settings, f'{{{ns}}}updateFields')
        uf.set(f'{{{ns}}}val', 'true')

    def _add_header_footer(self):
        """Add running header and footer with page numbers."""
        T = self.T
        section = self.doc.sections[0]

        # Header
        header_title = self.cfg.get("header_title", "")
        if header_title:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            _add_mixed_run(p, header_title, SANS, Pt(8), T["ink_faded"])
            # Add bottom border to header paragraph
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{T["border"]}" />'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add top border to footer paragraph
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="{T["border"]}" />'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Footer left: author
        footer_left = self.cfg.get("footer_left", self.cfg.get("author", ""))
        if footer_left:
            _add_mixed_run(p, footer_left + "    ", SANS, Pt(8), T["ink_faded"])

        # Page number field
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = T["accent"]
        fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fld_begin)
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instr)
        run3 = p.add_run()
        fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._element.append(fld_sep)
        run4 = p.add_run("1")
        run4.font.size = Pt(9)
        run4.font.color.rgb = T["accent"]
        run5 = p.add_run()
        fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._element.append(fld_end)

    def _add_heading(self, text, level):
        """Add heading with CJK support."""
        p = self.doc.add_heading(level=level)
        _add_inline_content(p, text, self.T, self.T["heading_font"],
                           Pt([0, 26, 18, 13][level]))
        return p

    def _add_paragraph(self, text):
        """Add body paragraph with inline markdown and CJK support."""
        p = self.doc.add_paragraph()
        _add_inline_content(p, text, self.T, self.T["body_font"], Pt(10.5))
        return p

    def _add_bullet(self, text, ordered=False, number=None):
        """Add bullet or numbered list item."""
        style = 'List Number' if ordered else 'List Bullet'
        p = self.doc.add_paragraph(style=style)
        _add_inline_content(p, text, self.T, self.T["body_font"], Pt(10.5))
        return p

    def _add_code_block(self, code):
        """Add code block with shaded background."""
        T = self.T
        p = self.doc.add_paragraph()
        # Shaded background via paragraph shading
        shd_hex = T["canvas_sec"].lstrip('#')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shd_hex}" w:val="clear"/>')
        p._element.get_or_add_pPr().append(shading)

        # Border
        border_hex = str(T["border"])
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{border_hex}" />'
            f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="{border_hex}" />'
            f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="{border_hex}" />'
            f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="{border_hex}" />'
            f'</w:pBdr>'
        )
        p._element.get_or_add_pPr().append(pBdr)

        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(14)

        # Add code text with font switching for CJK in code
        lines = code.split('\n')
        code_max = self.cfg.get("code_max_lines", 30)
        if len(lines) > code_max:
            lines = lines[:code_max - 1] + ['  // ... (truncated)']

        for idx, line in enumerate(lines):
            for is_cjk, seg in _split_mixed(line):
                run = p.add_run(seg)
                run.font.name = CJK if is_cjk else MONO
                run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x3D, 0x3D, 0x3A)
            if idx < len(lines) - 1:
                run = p.add_run('\n')
                run.font.size = Pt(8.5)

    def _add_table(self, lines):
        """Parse markdown table and add to document."""
        rows = []
        for l in lines:
            l = l.strip().strip('|')
            rows.append([c.strip() for c in l.split('|')])
        if len(rows) < 2:
            return
        header = rows[0]
        data = [r for r in rows[1:] if not all(set(c.strip()) <= set('-: ') for c in r)]
        if not data:
            return

        T = self.T
        nc = len(header)
        table = self.doc.add_table(rows=1 + len(data), cols=nc)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for ci, h in enumerate(header):
            cell = table.rows[0].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_mixed_run(p, h, SANS, Pt(9), RGBColor(0xFF, 0xFF, 0xFF), bold=True)
            # Accent background
            accent_hex = str(T["accent"])
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{accent_hex}" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

        # Data rows
        for ri, row in enumerate(data):
            for ci in range(nc):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                val = row[ci] if ci < len(row) else ""
                _add_inline_content(p, val, T, T["body_font"], Pt(9))
                # Alternating row shading
                if ri % 2 == 1:
                    shd_hex = T["canvas_sec"].lstrip('#')
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shd_hex}" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    def _add_image(self, src, alt=""):
        """Add image to document. Supports local paths and remote URLs."""
        img_path = None
        tmp_file = None
        try:
            if src.startswith(('http://', 'https://')):
                # Strip image processing query params for cleaner download
                clean_url = src.split('?')[0]
                suffix = Path(clean_url).suffix or '.jpg'
                tmp_file = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
                tmp_file.close()
                urllib.request.urlretrieve(clean_url, tmp_file.name)
                img_path = tmp_file.name
            else:
                # Resolve relative to input file directory
                base_dir = self.cfg.get("_input_dir", ".")
                candidate = os.path.join(base_dir, src)
                if os.path.isfile(candidate):
                    img_path = candidate
                elif os.path.isfile(src):
                    img_path = src

            if img_path and os.path.isfile(img_path):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.0))
                if alt:
                    cap = self.doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_mixed_run(cap, alt, self.T["body_font"], Pt(9), self.T["ink_faded"], italic=True)
            else:
                self._add_paragraph(f"[Image: {alt or src}]")
        except Exception as e:
            print(f"Warning: Could not embed image '{src}': {e}", file=sys.stderr)
            self._add_paragraph(f"[Image: {alt or src}]")
        finally:
            if tmp_file and os.path.exists(tmp_file.name):
                os.unlink(tmp_file.name)

    def _add_blockquote(self, text):
        """Add blockquote with left border."""
        T = self.T
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(10)

        accent_hex = str(T["accent"])
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{accent_hex}" />'
            f'</w:pBdr>'
        )
        p._element.get_or_add_pPr().append(pBdr)

        _add_inline_content(p, text, T, T["body_font"], Pt(10))

    @staticmethod
    def _preprocess_md(md):
        """Normalize markdown: strip YAML frontmatter, split merged headings."""
        # Strip YAML frontmatter
        if md.startswith('---'):
            end = md.find('\n---', 3)
            if end != -1:
                md = md[end + 4:]

        lines = md.split('\n')
        out = []
        in_code = False
        for line in lines:
            if line.strip().startswith('```'):
                in_code = not in_code
            if in_code:
                out.append(line)
                continue
            parts = re.split(r'(?<=[^#\s])\s*(?=#{1,3}\s)', line)
            if len(parts) > 1:
                for p in parts:
                    p = p.strip()
                    if p:
                        out.append(p)
            else:
                out.append(line)
        return '\n'.join(out)

    def parse_md(self, md):
        """Parse markdown and build document content."""
        md = self._preprocess_md(md)
        lines = md.split('\n')
        i = 0
        in_code = False
        code_buf = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Code blocks
            if stripped.startswith('```'):
                if in_code:
                    ct = '\n'.join(code_buf)
                    if ct.strip():
                        self._add_code_block(ct)
                    code_buf = []
                    in_code = False
                else:
                    in_code = True
                    code_buf = []
                i += 1
                continue
            if in_code:
                code_buf.append(line)
                i += 1
                continue

            # Skip frontmatter-like lines and empty
            if stripped in ('---', '\\newpage', '') or \
               stripped.startswith(('title:', 'subtitle:', 'author:', 'date:')):
                i += 1
                continue

            # H1
            if stripped.startswith('# ') and not stripped.startswith('## '):
                title = stripped.lstrip('#').strip()
                self._add_heading(title, 1)
                i += 1
                continue

            # H2
            if stripped.startswith('## ') and not stripped.startswith('### '):
                title = stripped[3:].strip()
                self._add_heading(title, 2)
                i += 1
                continue

            # H3
            if stripped.startswith('### '):
                title = stripped[4:].strip()
                self._add_heading(title, 3)
                i += 1
                continue

            # Tables
            if stripped.startswith('|'):
                tl = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    tl.append(lines[i])
                    i += 1
                self._add_table(tl)
                continue

            # Bullets
            if stripped.startswith('- ') or stripped.startswith('* '):
                self._add_bullet(stripped[2:].strip())
                i += 1
                continue

            # Numbered list
            m = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if m:
                self._add_bullet(m.group(2), ordered=True, number=m.group(1))
                i += 1
                continue

            # Images: ![alt](src)
            img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
            if img_match:
                alt, src = img_match.group(1), img_match.group(2)
                self._add_image(src, alt)
                i += 1
                continue

            # Blockquote
            if stripped.startswith('> '):
                self._add_blockquote(stripped[2:].strip())
                i += 1
                continue

            # Paragraph — join consecutive lines
            plines = []
            while i < len(lines):
                l = lines[i].strip()
                if not l or l.startswith('#') or l.startswith('```') or \
                   l.startswith('|') or l.startswith('- ') or l.startswith('* ') or \
                   l.startswith('> ') or re.match(r'^\d+\.\s', l):
                    break
                plines.append(l)
                i += 1
            if plines:
                merged = plines[0]
                for pl in plines[1:]:
                    if merged and pl and _is_cjk(merged[-1]) and _is_cjk(pl[0]):
                        merged += pl
                    else:
                        merged += ' ' + pl
                self._add_paragraph(merged)
            continue

    def build(self, md_text, output_path):
        """Build the DOCX file."""
        print("Building DOCX...")

        # Cover page
        if self.cfg.get("cover", True):
            self._add_cover()

        # TOC
        if self.cfg.get("toc", True):
            self._add_toc(md_text)

        # Watermark
        wm = self.cfg.get("watermark", "")
        if wm:
            self._add_watermark(wm)

        # Header/footer
        self._add_header_footer()

        # Parse and add content
        self.parse_md(md_text)

        # Auto-update fields (TOC) on open
        self._enable_auto_update_fields()

        # Save
        self.doc.save(output_path)
        size = os.path.getsize(output_path)
        print(f"Done! {output_path} ({size/1024:.1f} KB)")

# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description="md2docx — Markdown to Professional DOCX")
    parser.add_argument("--input", "-i", required=True, help="Input markdown file")
    parser.add_argument("--output", "-o", default="output.docx", help="Output DOCX path")
    parser.add_argument("--title", default="", help="Cover page title")
    parser.add_argument("--subtitle", default="", help="Cover page subtitle")
    parser.add_argument("--author", default="", help="Author name")
    parser.add_argument("--date", default=str(date.today()), help="Date string")
    parser.add_argument("--version", default="", help="Version string on cover")
    parser.add_argument("--watermark", default="", help="Watermark text (empty = none)")
    parser.add_argument("--theme", default="warm-academic", help="Theme name")
    parser.add_argument("--cover", default=True, type=lambda x: x.lower() != 'false', help="Generate cover page")
    parser.add_argument("--toc", default=True, type=lambda x: x.lower() != 'false', help="Generate TOC")
    parser.add_argument("--header-title", default="", help="Report title in header")
    parser.add_argument("--footer-left", default="", help="Brand/author in footer")
    parser.add_argument("--stats-line", default="", help="Stats line on cover")
    parser.add_argument("--stats-line2", default="", help="Second stats line on cover")
    parser.add_argument("--edition-line", default="", help="Edition line on cover")
    parser.add_argument("--code-max-lines", default=30, type=int, help="Max lines per code block")
    args = parser.parse_args()

    with open(args.input, encoding='utf-8') as f:
        md_text = f.read()

    # Extract title from first H1 if not provided
    title = args.title
    if not title:
        m = re.search(r'^# (.+)$', md_text, re.MULTILINE)
        title = m.group(1).strip() if m else "Document"

    theme = load_theme(args.theme)
    config = {
        "title": title,
        "subtitle": args.subtitle,
        "author": args.author,
        "date": args.date,
        "version": args.version,
        "watermark": args.watermark,
        "theme": theme,
        "cover": args.cover,
        "toc": args.toc,
        "header_title": args.header_title,
        "footer_left": args.footer_left or args.author,
        "stats_line": args.stats_line,
        "stats_line2": args.stats_line2,
        "edition_line": args.edition_line,
        "code_max_lines": args.code_max_lines,
        "_input_dir": os.path.dirname(os.path.abspath(args.input)),
    }

    builder = DocxBuilder(config)
    builder.build(md_text, args.output)

if __name__ == "__main__":
    main()
